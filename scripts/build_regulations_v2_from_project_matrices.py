from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.document import Document as DocumentObject
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.shared import Cm, Pt
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph

from skill_assessment.data.top20_position_kpis import TOP20_POSITION_KPI_ROWS
from skill_assessment.data.top20_position_skills import TOP20_POSITION_SKILL_ROWS


def _resolve_regulations_dirs() -> tuple[Path, Path]:
    """Ищет typical_infrastructure/docs/regulations/V1 вверх по дереву каталогов.

    Работает и для клона только skill_assessment рядом с Stage3HR, и для копии
    внутри Brotherly_hearts_training_tracks/SergeyK3/skill_assessment при общем корне Stage3HR.
    """
    here = Path(__file__).resolve()
    for base in here.parents:
        v1 = base / "typical_infrastructure" / "docs" / "regulations" / "V1"
        if v1.is_dir():
            return v1, v1.parent / "V2"
    raise FileNotFoundError(
        "Не найдена папка typical_infrastructure/docs/regulations/V1 "
        "(ожидается рядом с корнем рабочей копии Stage3HR)."
    )


SOURCE_DIR, TARGET_DIR = _resolve_regulations_dirs()
CURRENT_DATE_STR = "28.03.2026"
DOCUMENT_VERSION = "V2"
MATRIX_VERSION_LABEL = f"{DOCUMENT_VERSION} от {CURRENT_DATE_STR}"

SKILLS_INFO_TEXT = (
    "Информация по навыкам может устаревать. Актуальную версию смотрите в своих "
    "должностных документах или запросите в отделе кадров. Навыки измеряются по "
    "кейсам и руководителем."
)
SKILLS_AFTER_TEXT = "После обновления навыков следует детализировать целевой ориентир."
KPI_INFO_TEXT = (
    "Информация по KPI может устаревать. Актуальную версию смотрите в своих "
    "должностных документах или запросите в отделе кадров."
)
KPI_AFTER_TEXT = (
    "После обновления ключевых показателей следует детализировать целевой ориентир."
)

POSITION_CODE_ALIASES = {
    "DIRECTOR": "ADM_DIRECTOR",
    "SYSADMIN": "ADM_SYS_ADMIN",
    "SALES_MGR": "SALES_MANAGER",
}

SKILLS_BY_POSITION = {
    position_code: skills for position_code, _position_name, skills in TOP20_POSITION_SKILL_ROWS
}
KPIS_BY_POSITION = {
    position_code: kpis for position_code, _position_name, kpis in TOP20_POSITION_KPI_ROWS
}


def iter_block_items(parent):
    if isinstance(parent, DocumentObject):
        parent_elm = parent.element.body
    elif isinstance(parent, _Cell):
        parent_elm = parent._tc
    else:
        raise TypeError(f"Unsupported parent type: {type(parent)!r}")

    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)


def block_element(block):
    if isinstance(block, Paragraph):
        return block._p
    if isinstance(block, Table):
        return block._tbl
    raise TypeError(f"Unsupported block type: {type(block)!r}")


def insert_paragraph_before(block, text: str, style=None) -> Paragraph:
    new_p = OxmlElement("w:p")
    block_element(block).addprevious(new_p)
    paragraph = Paragraph(new_p, block._parent)
    if style is not None:
        paragraph.style = style
    if text:
        paragraph.add_run(text)
    compact_paragraph(paragraph)
    return paragraph


def insert_paragraph_after(block, text: str, style=None) -> Paragraph:
    new_p = OxmlElement("w:p")
    block_element(block).addnext(new_p)
    paragraph = Paragraph(new_p, block._parent)
    if style is not None:
        paragraph.style = style
    if text:
        paragraph.add_run(text)
    compact_paragraph(paragraph)
    return paragraph


def insert_table_before(doc: Document, block, rows: int, cols: int, style=None) -> Table:
    table = doc.add_table(rows=rows, cols=cols)
    if style is not None:
        table.style = style
    block_element(block).addprevious(table._tbl)
    return table


def remove_block(block) -> None:
    element = block_element(block)
    element.getparent().remove(element)


def increment_regulation_code(value: str) -> str:
    value = value.strip()
    return re.sub(r"_V(\d+)$", lambda m: f"_V{int(m.group(1)) + 1}", value)


def increment_version(value: str) -> str:
    value = value.strip()
    match = re.fullmatch(r"V(\d+)", value, flags=re.IGNORECASE)
    if not match:
        return value
    return f"V{int(match.group(1)) + 1}"


def normalize_position_code(raw_code: str) -> str:
    code = raw_code.strip()
    return POSITION_CODE_ALIASES.get(code, code)


def extract_position_code(doc: Document) -> str:
    meta = doc.tables[0]
    regulation_code = meta.cell(0, 1).text.strip()
    duty = meta.cell(2, 1).text.replace("\n", " ").strip()
    match = re.search(r"\(([A-Z][A-Z0-9_]*)\)", duty)
    if match:
        return normalize_position_code(match.group(1))
    match = re.match(r"REG_(.+)_V\d+$", regulation_code)
    if match:
        return normalize_position_code(match.group(1))
    raise ValueError(f"Не удалось определить position_code: {regulation_code!r}")


def renumber_following_sections(doc: Document) -> None:
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        match = re.match(r"^(\d+)\.\s+(.*)$", text)
        if not match:
            continue
        number = int(match.group(1))
        if number >= 7:
            paragraph.text = f"{number + 1}. {match.group(2)}"


def find_kpi_heading(doc: Document) -> Paragraph:
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if re.match(r"^\d+\.\s+", text) and "KPI" in text.upper():
            return paragraph
    raise ValueError("Не найден заголовок KPI-раздела")


def find_kpi_table(doc: Document) -> Table:
    for table in doc.tables:
        if not table.rows:
            continue
        header = [cell.text.replace("\n", " ").strip().lower() for cell in table.rows[0].cells]
        joined = " ".join(header)
        if "показатель" in joined and "как измеряется" in joined:
            return table
    raise ValueError("Не найдена KPI-таблица")


def find_existing_kpi_intro(doc: Document, kpi_heading: Paragraph, kpi_table: Table) -> Paragraph | None:
    after_heading = False
    for block in iter_block_items(doc):
        if block is kpi_heading:
            after_heading = True
            continue
        if not after_heading:
            continue
        if block is kpi_table:
            return None
        if isinstance(block, Paragraph) and block.text.strip():
            return block
    return None


def clear_extra_meta_fields(doc: Document) -> None:
    meta = doc.tables[0]
    meta.cell(0, 1).text = increment_regulation_code(meta.cell(0, 1).text)
    meta.cell(0, 3).text = increment_version(meta.cell(0, 3).text)
    meta.cell(2, 3).text = CURRENT_DATE_STR
    clear_cell(meta.cell(3, 2))
    clear_cell(meta.cell(3, 3))


def format_meta_table(doc: Document) -> None:
    meta = doc.tables[0]
    set_table_widths(meta, [Cm(4.0), Cm(8.8), Cm(2.0), Cm(2.5)])


def remove_duplicate_goal_block(doc: Document) -> None:
    if len(doc.tables) < 2:
        return
    duplicate_block = doc.tables[1]
    if len(duplicate_block.rows) == 1 and len(duplicate_block.columns) == 1:
        remove_block(duplicate_block)


def clear_table_rows_except_header(table: Table) -> None:
    for row in list(table.rows)[1:]:
        row._tr.getparent().remove(row._tr)


def compact_paragraph(paragraph: Paragraph) -> None:
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(0)


def compact_table_paragraphs(table: Table) -> None:
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                compact_paragraph(paragraph)


def clear_cell(cell: _Cell) -> None:
    cell.text = ""
    for paragraph in cell.paragraphs:
        compact_paragraph(paragraph)


def set_cell_width(cell: _Cell, width) -> None:
    cell.width = width
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.tcW
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:type"), "dxa")
    tc_w.set(qn("w:w"), str(width.twips))


def set_table_widths(table: Table, widths: list) -> None:
    table.autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths):
            set_cell_width(row.cells[idx], width)
    compact_table_paragraphs(table)


def remove_empty_paragraphs(doc: Document) -> None:
    for paragraph in list(doc.paragraphs):
        if paragraph.text.strip():
            continue
        if paragraph._p.xpath(".//w:drawing"):
            continue
        remove_block(paragraph)


def normalize_heading_spacing(doc: Document) -> None:
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if re.match(r"^\d+\.\s+", text):
            compact_paragraph(paragraph)
            paragraph.paragraph_format.space_before = Pt(12)


def format_target(value, unit: str) -> str:
    if value is None:
        return "Требует детализации"
    if isinstance(value, float) and value.is_integer():
        number = str(int(value))
    else:
        number = str(value).replace(".", ",")
    if unit == "%":
        return f"{number}%"
    return f"{number} {unit}".strip()


def period_type_label(period_type: str) -> str:
    return {
        "month": "месяц",
        "quarter": "квартал",
        "year": "год",
    }.get(period_type, period_type)


def fill_skills_table(table: Table, skills: list[str]) -> None:
    headers = ("№", "Навык", "Приоритет", "Версия")
    for idx, header in enumerate(headers):
        table.cell(0, idx).text = header
    for rank, skill in enumerate(skills, start=1):
        row = table.add_row().cells
        row[0].text = str(rank)
        row[1].text = skill
        row[2].text = str(rank)
        row[3].text = MATRIX_VERSION_LABEL
    set_table_widths(table, [Cm(0.9), Cm(9.6), Cm(1.7), Cm(3.0)])


def refill_kpi_table(table: Table, kpis: list[tuple[str, str, str, str, float | None]]) -> None:
    while len(table.columns) < 5:
        table.add_column(Cm(2.4))

    headers = ("№", "Показатель", "Как измеряется", "Код и версия", "Целевой ориентир")
    for idx, header in enumerate(headers):
        table.cell(0, idx).text = header

    clear_table_rows_except_header(table)

    for rank, (kpi_code, title_ru, unit, period_type, default_target) in enumerate(kpis, start=1):
        row = table.add_row().cells
        row[0].text = str(rank)
        row[1].text = title_ru
        row[2].text = f"Единица: {unit}; период: {period_type_label(period_type)}"
        row[3].text = f"{kpi_code}\n{MATRIX_VERSION_LABEL}"
        row[4].text = format_target(default_target, unit)
    set_table_widths(table, [Cm(0.8), Cm(4.2), Cm(6.3), Cm(3.7), Cm(2.4)])


def add_skills_section(doc: Document, kpi_heading: Paragraph, skills: list[str], table_style, heading_style, body_style) -> None:
    insert_paragraph_before(kpi_heading, "7. Ключевые навыки", style=heading_style)
    insert_paragraph_before(kpi_heading, SKILLS_INFO_TEXT, style=body_style)
    skills_table = insert_table_before(doc, kpi_heading, rows=1, cols=4, style=table_style)
    fill_skills_table(skills_table, skills)
    insert_paragraph_before(kpi_heading, SKILLS_AFTER_TEXT, style=body_style)


def update_document(path: Path) -> tuple[Path, str]:
    doc = Document(str(path))
    position_code = extract_position_code(doc)
    skills = SKILLS_BY_POSITION.get(position_code)
    kpis = KPIS_BY_POSITION.get(position_code)

    if not skills:
        raise KeyError(f"Не найдены навыки для {position_code}")
    if not kpis:
        raise KeyError(f"Не найдены KPI для {position_code}")

    clear_extra_meta_fields(doc)
    format_meta_table(doc)
    remove_duplicate_goal_block(doc)
    renumber_following_sections(doc)

    kpi_heading = find_kpi_heading(doc)
    kpi_table = find_kpi_table(doc)
    kpi_intro = find_existing_kpi_intro(doc, kpi_heading, kpi_table)

    heading_style = kpi_heading.style
    body_style = kpi_intro.style if kpi_intro is not None else doc.paragraphs[0].style
    table_style = kpi_table.style

    add_skills_section(doc, kpi_heading, skills, table_style, heading_style, body_style)
    insert_paragraph_before(kpi_table, KPI_INFO_TEXT, style=body_style)
    refill_kpi_table(kpi_table, kpis)
    insert_paragraph_after(kpi_table, KPI_AFTER_TEXT, style=body_style)
    normalize_heading_spacing(doc)
    remove_empty_paragraphs(doc)
    for table in doc.tables:
        compact_table_paragraphs(table)

    target_path = TARGET_DIR / path.name
    doc.save(str(target_path))
    return target_path, position_code


def main() -> None:
    TARGET_DIR.mkdir(parents=True, exist_ok=True)

    processed = []
    for path in sorted(SOURCE_DIR.glob("*.docx")):
        processed.append(update_document(path))

    print(f"Updated {len(processed)} regulations into {TARGET_DIR}")
    for target_path, position_code in processed:
        print(f"- {position_code}: {target_path.name}")


if __name__ == "__main__":
    main()
